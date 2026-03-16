from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

# ---- Вспомогательные функции ----

def add_run_tnr(paragraph, text, bold=False, italic=False, size=14):
    run = paragraph.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),  'Times New Roman')
    rFonts.set(qn('w:hAnsi'),  'Times New Roman')
    rFonts.set(qn('w:cs'),     'Times New Roman')
    rPr.insert(0, rFonts)
    return run

def set_line_spacing(paragraph, spacing=1.5):
    pPr = paragraph._p.get_or_add_pPr()
    lnSpc = OxmlElement('w:spacing')
    if spacing == 1.5:
        lnSpc.set(qn('w:line'), '360')
    else:
        lnSpc.set(qn('w:line'), '240')
    lnSpc.set(qn('w:lineRule'), 'auto')
    pPr.append(lnSpc)

def add_paragraph(document, text='', bold=False, italic=False, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0):
    p = document.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    set_line_spacing(p, spacing)
    if text:
        add_run_tnr(p, text, bold=bold, italic=italic, size=size)
    return p

def body_paragraph(document, text, first_indent=True):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.first_line_indent = Cm(1.25) if first_indent else Cm(0)
    set_line_spacing(p, 1.5)
    add_run_tnr(p, text)
    return p

def heading_paragraph(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    pf.first_line_indent = Cm(0)
    set_line_spacing(p, 1.5)
    add_run_tnr(p, text, bold=True)
    return p

def add_page_break(document):
    p = document.add_paragraph()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p._p.append(br)

def code_paragraph(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.first_line_indent = Cm(0)
    set_line_spacing(p, 1.0)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),  'Courier New')
    rFonts.set(qn('w:hAnsi'),  'Courier New')
    rFonts.set(qn('w:cs'),     'Courier New')
    rPr.insert(0, rFonts)
    return p

def step_paragraph(document, text):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    p.paragraph_format.space_before = Pt(6)
    set_line_spacing(p, 1.5)
    add_run_tnr(p, text, bold=True)
    return p

# ======================================================
#  ТИТУЛЬНЫЙ ЛИСТ
# ======================================================
top_lines = [
    'МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ',
    'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ '
    'ВЫСШЕГО ОБРАЗОВАНИЯ «НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ ТЕХНОЛОГИЧЕСКИЙ '
    'УНИВЕРСИТЕТ «МИСИС»',
    'ИНСТИТУТ КОМПЬЮТЕРНЫХ НАУК (ИКН)',
    'КАФЕДРА АВТОМАТИЗИРОВАННЫХ СИСТЕМ УПРАВЛЕНИЯ (АСУ)',
]
for line in top_lines:
    add_paragraph(doc, line, bold=True, size=14)

for _ in range(6):
    add_paragraph(doc, '')

add_paragraph(doc,
    'Курс «Клиент-серверные приложения и сетевые технологии»')
add_paragraph(doc, 'Практическая работа №\u00a03')
add_paragraph(doc, '«Добавление динамики с JavaScript»', bold=True)

for _ in range(5):
    add_paragraph(doc, '')

def right_p(text, underline=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    set_line_spacing(p, 1.0)
    run = p.add_run(text)
    run.underline = underline
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p

right_p('Выполнил: студент группы БИВТ-24-2')
right_p('Горбачёв В.\u00a0А.', underline=True)
right_p('Проверили:           Абросимов Н.\u00a0А.')
right_p('                            Шелудяков П.\u00a0А.')

for _ in range(4):
    add_paragraph(doc, '')

add_paragraph(doc, 'Москва, 2026')

add_page_break(doc)

# ======================================================
#  ЦЕЛЬ РАБОТЫ
# ======================================================
heading_paragraph(doc, 'Цель работы')

body_paragraph(doc,
    'Изучить первичные основы языка JavaScript и получить навыки '
    'в практической реализации его возможностей на страницах сайта '
    '«Каталог смартфонов». Реализовать фильтрацию товаров, корзину '
    'с подсчётом стоимости и оплату.')

# ======================================================
#  ХОД РАБОТЫ
# ======================================================
heading_paragraph(doc, 'Ход работы')

body_paragraph(doc,
    'Для выполнения работы использовался текстовый редактор '
    'Visual Studio Code и браузер Google Chrome. В качестве основы '
    'взят сайт «Каталог смартфонов», разработанный в практических '
    'работах №\u00a01–2.')

# --- Шаг 1 ---
step_paragraph(doc,
    'Шаг 1. Создание внешнего файла script.js и подключение к страницам.')

body_paragraph(doc,
    'Создан файл script.js, содержащий весь JavaScript-код сайта. '
    'Файл подключён ко всем HTML-страницам перед закрывающим тегом '
    '</body> с помощью тега <script>:')

code_paragraph(doc, '<script src="script.js"></script>')

# --- Шаг 2 ---
step_paragraph(doc,
    'Шаг 2. Создание глобальной переменной корзины и массива товаров.')

body_paragraph(doc,
    'Объявлена глобальная переменная cart — пустой массив для хранения '
    'товаров корзины. Также создан массив products с данными о товарах '
    '(название, цена, категория):')

code_paragraph(doc, 'let cart = [];')
code_paragraph(doc, '')
code_paragraph(doc, 'const products = [')
code_paragraph(doc, '    { id: 1, name: "iPhone 15 Pro", price: 129990, category: "apple" },')
code_paragraph(doc, '    { id: 2, name: "Samsung Galaxy S24 Ultra", price: 109990, category: "samsung" },')
code_paragraph(doc, '    { id: 3, name: "Google Pixel 8 Pro", price: 84990, category: "google" }')
code_paragraph(doc, '];')

# --- Шаг 3 ---
step_paragraph(doc,
    'Шаг 3. Реализация стрелочных функций для работы с корзиной.')

body_paragraph(doc,
    'Все основные функции реализованы в виде стрелочных функций (arrow '
    'functions). Функция calculateTotal() выполняет подсчёт общей суммы '
    'товаров в корзине с помощью метода forEach:')

code_paragraph(doc, 'const calculateTotal = () => {')
code_paragraph(doc, '    let total = 0;')
code_paragraph(doc, '    cart.forEach(item => total += item.price);')
code_paragraph(doc, '    return total;')
code_paragraph(doc, '};')

body_paragraph(doc,
    'Функция addToCart() добавляет товар в массив корзины и вызывает '
    'перерисовку:')

code_paragraph(doc, 'const addToCart = (product) => {')
code_paragraph(doc, '    cart.push(product);')
code_paragraph(doc, '    renderCart();')
code_paragraph(doc, '};')

body_paragraph(doc,
    'Функция removeFromCart() удаляет товар из корзины по индексу '
    'с помощью метода splice:')

code_paragraph(doc, 'const removeFromCart = (index) => {')
code_paragraph(doc, '    cart.splice(index, 1);')
code_paragraph(doc, '    renderCart();')
code_paragraph(doc, '};')

body_paragraph(doc,
    'Функция clearCart() полностью очищает массив корзины:')

code_paragraph(doc, 'const clearCart = () => {')
code_paragraph(doc, '    cart = [];')
code_paragraph(doc, '    renderCart();')
code_paragraph(doc, '};')

# --- Шаг 4 ---
step_paragraph(doc,
    'Шаг 4. Реализация кнопки «Оплатить».')

body_paragraph(doc,
    'Функция pay() проверяет, пуста ли корзина. Если корзина пуста — '
    'выводится всплывающее окно alert("Корзина пуста!"). При успешной '
    'оплате выводится сообщение с суммой и корзина очищается:')

code_paragraph(doc, 'const pay = () => {')
code_paragraph(doc, '    if (cart.length === 0) {')
code_paragraph(doc, '        alert("Корзина пуста!");')
code_paragraph(doc, '        return;')
code_paragraph(doc, '    }')
code_paragraph(doc, '    alert("Оплата прошла успешно! Сумма: " + calculateTotal() + " руб.");')
code_paragraph(doc, '    clearCart();')
code_paragraph(doc, '};')

# --- Шаг 5 ---
step_paragraph(doc,
    'Шаг 5. Отрисовка корзины (функция renderCart).')

body_paragraph(doc,
    'Функция renderCart() выполняет следующие действия: очищает текущий '
    'список корзины (innerHTML = ""); заново отрисовывает каждый товар '
    'с кнопкой «Удалить»; пересчитывает общую сумму через calculateTotal(); '
    'обновляет текст итоговой суммы через свойство textContent.')

code_paragraph(doc, 'const renderCart = () => {')
code_paragraph(doc, '    const cartList = document.querySelector("#cart-list");')
code_paragraph(doc, '    const cartTotal = document.querySelector("#cart-total");')
code_paragraph(doc, '    cartList.innerHTML = "";')
code_paragraph(doc, '    cart.forEach((item, index) => {')
code_paragraph(doc, '        // создание элементов для каждого товара')
code_paragraph(doc, '        // с кнопкой "Удалить"')
code_paragraph(doc, '    });')
code_paragraph(doc, '    cartTotal.textContent = "Итого: " + calculateTotal() + " руб.";')
code_paragraph(doc, '};')

# --- Шаг 6 ---
step_paragraph(doc,
    'Шаг 6. Реализация фильтра товаров в каталоге.')

body_paragraph(doc,
    'На странице каталога добавлен элемент <select> с вариантами '
    'фильтрации по бренду: «Все», «Apple», «Samsung», «Google». '
    'Каждая карточка товара содержит атрибут data-category с названием '
    'бренда. Функция filterProducts() перебирает все карточки и '
    'скрывает/показывает их в зависимости от выбранной категории:')

code_paragraph(doc, 'const filterProducts = (category) => {')
code_paragraph(doc, '    const cards = document.querySelectorAll(".product-card");')
code_paragraph(doc, '    cards.forEach(card => {')
code_paragraph(doc, '        const cardCategory = card.dataset.category;')
code_paragraph(doc, '        if (category === "all" || cardCategory === category) {')
code_paragraph(doc, '            card.style.display = "inline-block";')
code_paragraph(doc, '        } else {')
code_paragraph(doc, '            card.style.display = "none";')
code_paragraph(doc, '        }')
code_paragraph(doc, '    });')
code_paragraph(doc, '};')

# --- Шаг 7 ---
step_paragraph(doc,
    'Шаг 7. Обработка событий.')

body_paragraph(doc,
    'Все обработчики событий назначены с помощью метода addEventListener(). '
    'Для кнопок «Добавить в корзину» данные о товаре (название и цена) '
    'извлекаются из data-атрибутов кнопки с помощью свойства dataset. '
    'Цена преобразуется из строки в число с помощью Number():')

code_paragraph(doc, 'button.addEventListener("click", () => {')
code_paragraph(doc, '    const name = button.dataset.name;')
code_paragraph(doc, '    const price = Number(button.dataset.price);')
code_paragraph(doc, '    addToCart({ name, price });')
code_paragraph(doc, '});')

body_paragraph(doc,
    'Для фильтра используется событие "change" элемента <select>:')

code_paragraph(doc, 'filterSelect.addEventListener("change", () => {')
code_paragraph(doc, '    filterProducts(filterSelect.value);')
code_paragraph(doc, '});')

# --- Рисунки ---
body_paragraph(doc,
    'На рисунке\u00a01 представлен внешний вид страницы каталога '
    'с фильтром и корзиной.')

p_fig = doc.add_paragraph()
p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_fig, 1.0)
add_run_tnr(p_fig, '[Рисунок 1 — Каталог с фильтром и корзиной]', italic=True)

p_cap = doc.add_paragraph()
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_cap, 1.0)
add_run_tnr(p_cap, 'Рисунок 1 — Каталог с фильтром и корзиной')

body_paragraph(doc,
    'На рисунке\u00a02 показана работа фильтра — отображение '
    'только товаров бренда Samsung.')

p_fig2 = doc.add_paragraph()
p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig2.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_fig2, 1.0)
add_run_tnr(p_fig2, '[Рисунок 2 — Фильтрация по бренду Samsung]', italic=True)

p_cap2 = doc.add_paragraph()
p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap2.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_cap2, 1.0)
add_run_tnr(p_cap2, 'Рисунок 2 — Фильтрация по бренду Samsung')

body_paragraph(doc,
    'На рисунке\u00a03 показана корзина с добавленными товарами '
    'и итоговой суммой.')

p_fig3 = doc.add_paragraph()
p_fig3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig3.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_fig3, 1.0)
add_run_tnr(p_fig3, '[Рисунок 3 — Корзина с товарами]', italic=True)

p_cap3 = doc.add_paragraph()
p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap3.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_cap3, 1.0)
add_run_tnr(p_cap3, 'Рисунок 3 — Корзина с товарами')

body_paragraph(doc,
    'На рисунке\u00a04 показано всплывающее окно при попытке '
    'оплаты пустой корзины и сообщение об успешной оплате.')

p_fig4 = doc.add_paragraph()
p_fig4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig4.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_fig4, 1.0)
add_run_tnr(p_fig4, '[Рисунок 4 — Всплывающие окна оплаты]', italic=True)

p_cap4 = doc.add_paragraph()
p_cap4.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap4.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_cap4, 1.0)
add_run_tnr(p_cap4, 'Рисунок 4 — Всплывающие окна оплаты')

# --- Таблица ---
body_paragraph(doc,
    'В таблице\u00a01 приведены основные конструкции JavaScript, '
    'использованные в работе.')

p_tbl_cap = doc.add_paragraph()
p_tbl_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_tbl_cap.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_tbl_cap, 1.0)
add_run_tnr(p_tbl_cap, 'Таблица 1 — Основные конструкции JavaScript')

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'

hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].clear()
hdr[1].paragraphs[0].clear()
add_run_tnr(hdr[0].paragraphs[0], 'Конструкция', bold=True)
add_run_tnr(hdr[1].paragraphs[0], 'Назначение', bold=True)
hdr[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

items = [
    ('let / const',           'Объявление переменных и констант'),
    ('[]  (массив)',          'Хранение списка значений (корзина товаров)'),
    ('push() / splice()',     'Добавление и удаление элементов массива'),
    ('forEach()',             'Перебор элементов массива'),
    ('() => {}',              'Стрелочная функция (arrow function)'),
    ('querySelector()',       'Получение элемента DOM по CSS-селектору'),
    ('querySelectorAll()',    'Получение всех элементов DOM по селектору'),
    ('addEventListener()',    'Назначение обработчика событий'),
    ('dataset',              'Доступ к data-атрибутам HTML-элемента'),
    ('style.display',        'Управление видимостью элемента (none/block)'),
    ('textContent',          'Изменение текстового содержимого элемента'),
    ('innerHTML',            'Изменение HTML-содержимого элемента'),
    ('alert()',              'Вывод всплывающего окна с сообщением'),
    ('Number()',             'Преобразование строки в число'),
]
for name, desc in items:
    row = tbl.add_row()
    row.cells[0].paragraphs[0].clear()
    row.cells[1].paragraphs[0].clear()
    add_run_tnr(row.cells[0].paragraphs[0], name)
    add_run_tnr(row.cells[1].paragraphs[0], desc)

# ======================================================
#  ВЫВОД
# ======================================================
heading_paragraph(doc, 'Вывод по ходу работы')

body_paragraph(doc,
    'В ходе выполнения практической работы были изучены основы языка '
    'JavaScript и его взаимодействие с HTML-документом через DOM '
    '(Document Object Model).')

body_paragraph(doc,
    'Были реализованы следующие функциональные возможности: фильтрация '
    'товаров в каталоге по категории (бренду) с помощью элемента <select> '
    'и обработчика события "change"; корзина товаров с возможностью '
    'добавления, удаления отдельных товаров и полной очистки; подсчёт '
    'общей стоимости товаров в корзине; кнопка «Оплатить» с проверкой '
    'пустой корзины и выводом всплывающих окон через alert().')

body_paragraph(doc,
    'В коде использованы стрелочные функции (arrow functions), '
    'методы работы с массивами (push, splice, forEach), обработка '
    'событий через addEventListener(), доступ к data-атрибутам '
    'через свойство dataset. Данные о товарах хранятся в глобальном '
    'массиве products, а корзина — в массиве cart. '
    'Цель работы достигнута в полном объёме.')

# ======================================================
#  СОХРАНЕНИЕ
# ======================================================
import os
out_dir = os.path.dirname(os.path.abspath(__file__))
out = os.path.join(out_dir,
    'ПР-3_БИВТ-24-2_Горбачёв_В_А_JavaScript.docx')
doc.save(out)
print('Saved:', out)
