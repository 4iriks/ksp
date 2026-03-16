from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ---- Параметры страницы: все поля 2 см ----
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

# ---- Вспомогательные функции ----

def add_paragraph(document, text='', bold=False, italic=False, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0,
                  first_line_indent=None, space_before=0, space_after=0):
    p = document.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    # межстрочный интервал
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    lnSpc = OxmlElement('w:spacing')
    if spacing == 1.0:
        lnSpc.set(qn('w:line'), '240')
        lnSpc.set(qn('w:lineRule'), 'auto')
    elif spacing == 1.5:
        lnSpc.set(qn('w:line'), '360')
        lnSpc.set(qn('w:lineRule'), 'auto')
    pPr.append(lnSpc)

    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'),    'Times New Roman')
        rFonts.set(qn('w:hAnsi'),    'Times New Roman')
        rFonts.set(qn('w:cs'),       'Times New Roman')
        rPr.insert(0, rFonts)
    return p

def add_run_tnr(paragraph, text, bold=False, italic=False, size=14):
    run = paragraph.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),  'Times New Roman')
    rFonts.set(qn('w:hAnsi'),  'Times New Roman')
    rFonts.set(qn('w:cs'),     'Times New Roman')
    rPr.insert(0, rFonts)
    return run

def set_line_spacing(paragraph, spacing=1.5):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    pPr = paragraph._p.get_or_add_pPr()
    lnSpc = OxmlElement('w:spacing')
    if spacing == 1.5:
        lnSpc.set(qn('w:line'), '360')
    else:
        lnSpc.set(qn('w:line'), '240')
    lnSpc.set(qn('w:lineRule'), 'auto')
    pPr.append(lnSpc)

def body_paragraph(document, text, first_indent=True):
    """Абзац основного текста: TNR 14pt, 1.5, выравн. по ширине, красная строка 1.25 см"""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    if first_indent:
        pf.first_line_indent = Cm(1.25)
    else:
        pf.first_line_indent = Cm(0)
    set_line_spacing(p, 1.5)
    add_run_tnr(p, text)
    return p

def heading_paragraph(document, text, level=1):
    """Заголовок раздела по ГОСТ: по центру, жирный, TNR 14pt"""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    pf.first_line_indent = Cm(0)
    set_line_spacing(p, 1.5)
    add_run_tnr(p, text, bold=True)
    return p

def add_page_break(document):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    p = document.add_paragraph()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p._p.append(br)

# ======================================================
#  ТИТУЛЬНЫЙ ЛИСТ  (формат из образца Фазлыева)
# ======================================================

# Всё шапочное — CAPS + BOLD + CENTER
top_lines = [
    'МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ',
    'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ '
    'ВЫСШЕГО ОБРАЗОВАНИЯ «НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ ТЕХНОЛОГИЧЕСКИЙ '
    'УНИВЕРСИТЕТ «МИСИС»',
    'ИНСТИТУТ КОМПЬЮТЕРНЫХ НАУК (ИКН)',
    'КАФЕДРА АВТОМАТИЗИРОВАННЫХ СИСТЕМ УПРАВЛЕНИЯ (АСУ)',
]
for line in top_lines:
    add_paragraph(doc, line, bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

# Отступ перед серединой
for _ in range(6):
    add_paragraph(doc, '', size=14, spacing=1.0)

# Курс — обычный, по центру
add_paragraph(doc,
    'Курс «Клиент-серверные приложения и сетевые технологии»',
    bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

# Вид работы — обычный
add_paragraph(doc, 'Практическая работа №\u00a01',
    bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

# Название — BOLD
add_paragraph(doc, '«Верстка сайта с HTML»',
    bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

# Отступ перед блоком «Выполнил»
for _ in range(5):
    add_paragraph(doc, '', size=14, spacing=1.0)

# Блок Выполнил / Проверили — выровнен вправо
def right_p(text, underline=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    set_line_spacing(p, 1.0)
    run = p.add_run(text)
    run.bold      = bold
    run.underline = underline
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p

right_p('Выполнил: студент группы БИВТ-24-2')
right_p('Горбачёв В.\u00a0А.', underline=True)
right_p('Проверили:           Абросимов Н.\u00a0А.')
right_p('                            Шелудяков П.\u00a0А.')

# Отступ перед городом
for _ in range(4):
    add_paragraph(doc, '', size=14, spacing=1.0)

# Город и год
add_paragraph(doc, 'Москва, 2026',
    bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

# ======================================================
#  РАЗРЫВ СТРАНИЦЫ — конец титульника
# ======================================================
add_page_break(doc)

# ======================================================
#  ЦЕЛЬ РАБОТЫ
# ======================================================
heading_paragraph(doc, 'Цель работы')

body_paragraph(doc,
    'На основе изучения языка разметки текстов HyperText Markup Language '
    '(HTML) разработать примитивный Web-сайт по тематике «Каталог '
    'смартфонов» из нескольких взаимосвязанных страниц.')

# ======================================================
#  ХОД РАБОТЫ
# ======================================================
heading_paragraph(doc, 'Ход работы')

body_paragraph(doc,
    'Для выполнения работы использовался текстовый редактор '
    'Visual Studio Code (VS Code) и браузер Google Chrome.')

body_paragraph(doc,
    'Была определена тема сайта — «Каталог смартфонов». Сайт содержит '
    'главную страницу, страницу каталога и три страницы отдельных товаров.')

body_paragraph(doc,
    'Структура сайта:')

# Список файлов
items = [
    'index.html — главная страница;',
    'catalog.html — страница каталога смартфонов;',
    'item_iphone.html — страница товара iPhone 15 Pro;',
    'item_samsung.html — страница товара Samsung Galaxy S24 Ultra;',
    'item_pixel.html — страница товара Google Pixel 8 Pro;',
    'images/ — папка с изображениями товаров.',
]
for item in items:
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.left_indent = Cm(1.25)
    set_line_spacing(p, 1.5)
    add_run_tnr(p, item)

# Шаг 1
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
p.paragraph_format.space_before = Pt(6)
set_line_spacing(p, 1.5)
add_run_tnr(p, 'Шаг 1. Создание главной страницы index.html.', bold=True)

body_paragraph(doc,
    'На главной странице размещены: навигационное меню в виде маркированного '
    'списка со ссылками на разделы сайта; горизонтальная черта (<hr>), '
    'отделяющая меню от основного содержимого; название сайта, выделенное '
    'тегом <h1>; приветственное сообщение в тегах <p>; горизонтальная черта '
    'перед подвалом; подвал со знаком авторского права (<small>&amp;copy;</small>) '
    'и текстом «Все права защищены». Меню и подвал присутствуют на каждой '
    'странице сайта.')

body_paragraph(doc,
    'На рисунке\u00a01 приведён внешний вид главной страницы сайта в браузере.')

p_fig = doc.add_paragraph()
p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_fig, 1.0)
add_run_tnr(p_fig, '[Рисунок 1 — Главная страница index.html]', italic=True)

p_cap = doc.add_paragraph()
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_cap, 1.0)
add_run_tnr(p_cap, 'Рисунок 1 — Главная страница index.html')

# Шаг 2
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
p.paragraph_format.space_before = Pt(6)
set_line_spacing(p, 1.5)
add_run_tnr(p, 'Шаг 2. Создание страницы каталога catalog.html.', bold=True)

body_paragraph(doc,
    'Страница каталога содержит: меню сайта; горизонтальную черту; заголовок '
    '«Каталог» (<h2>); уменьшенные изображения товаров размером 150×150 '
    'пикселей с атрибутами width и height тега <img>; ссылки под каждой '
    'картинкой для перехода на страницу подробного описания товара; подвал.')

body_paragraph(doc,
    'На рисунке\u00a02 приведён внешний вид страницы каталога.')

p_fig2 = doc.add_paragraph()
p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig2.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_fig2, 1.0)
add_run_tnr(p_fig2, '[Рисунок 2 — Страница каталога catalog.html]', italic=True)

p_cap2 = doc.add_paragraph()
p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap2.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_cap2, 1.0)
add_run_tnr(p_cap2, 'Рисунок 2 — Страница каталога catalog.html')

# Шаг 3
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
p.paragraph_format.space_before = Pt(6)
set_line_spacing(p, 1.5)
add_run_tnr(p, 'Шаг 3. Создание страниц элементов каталога.', bold=True)

body_paragraph(doc,
    'Для каждого из трёх товаров создана отдельная HTML-страница. На каждой '
    'странице присутствуют: меню; горизонтальная черта; название товара (<h2>); '
    'кликабельное изображение товара — при нажатии картинка открывается в '
    'полном размере в новой вкладке браузера (атрибут target="_blank" тега <a>); '
    'заголовок «Описание» (<h3>) с текстом описания; заголовок «Характеристики» '
    '(<h3>) с маркированным списком (<ul>, <li>) технических характеристик '
    'устройства; подвал.')

body_paragraph(doc,
    'На рисунке\u00a03 показана страница товара iPhone\u00a015\u00a0Pro '
    'в браузере.')

p_fig3 = doc.add_paragraph()
p_fig3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig3.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_fig3, 1.0)
add_run_tnr(p_fig3, '[Рисунок 3 — Страница товара item_iphone.html]', italic=True)

p_cap3 = doc.add_paragraph()
p_cap3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap3.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_cap3, 1.0)
add_run_tnr(p_cap3, 'Рисунок 3 — Страница товара item_iphone.html')

body_paragraph(doc,
    'Аналогичным образом оформлены страницы item_samsung.html '
    '(Samsung Galaxy S24 Ultra) и item_pixel.html (Google Pixel 8 Pro), '
    'представленные на рисунках\u00a04 и\u00a05 соответственно.')

for n, name in [(4, 'item_samsung.html'), (5, 'item_pixel.html')]:
    pf = doc.add_paragraph()
    pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf.paragraph_format.first_line_indent = Cm(0)
    set_line_spacing(pf, 1.0)
    add_run_tnr(pf, f'[Рисунок {n} — Страница товара {name}]', italic=True)

    pc = doc.add_paragraph()
    pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pc.paragraph_format.first_line_indent = Cm(0)
    set_line_spacing(pc, 1.0)
    add_run_tnr(pc, f'Рисунок {n} — Страница товара {name}')

# Основные теги — таблица (подпись ВЫШЕ таблицы по ГОСТ)
body_paragraph(doc,
    'В таблице\u00a01 приведены основные HTML-теги, применённые в работе.')

p_tbl_cap = doc.add_paragraph()
p_tbl_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_tbl_cap.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_tbl_cap, 1.0)
add_run_tnr(p_tbl_cap, 'Таблица 1 — Основные HTML-теги, использованные в работе')

tbl2 = doc.add_table(rows=1, cols=2)
tbl2.style = 'Table Grid'

hdr = tbl2.rows[0].cells
hdr[0].paragraphs[0].clear()
hdr[1].paragraphs[0].clear()
add_run_tnr(hdr[0].paragraphs[0], 'Тег', bold=True)
add_run_tnr(hdr[1].paragraphs[0], 'Назначение', bold=True)
hdr[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

tags = [
    ('<h1>, <h2>, <h3>', 'Заголовки первого, второго и третьего уровней'),
    ('<p>',              'Параграф (абзац текста)'),
    ('<hr>',             'Горизонтальная разделительная черта'),
    ('<ul>, <li>',       'Маркированный список и его элементы'),
    ('<a href="">',      'Гиперссылка (относительная и абсолютная)'),
    ('<img>',            'Встроенное изображение'),
    ('<b>, <i>',         'Жирное и курсивное начертание текста'),
    ('<small>',          'Уменьшенный размер шрифта'),
]
for tag, desc in tags:
    row = tbl2.add_row()
    row.cells[0].paragraphs[0].clear()
    row.cells[1].paragraphs[0].clear()
    add_run_tnr(row.cells[0].paragraphs[0], tag)
    add_run_tnr(row.cells[1].paragraphs[0], desc)

# ======================================================
#  ВЫВОД
# ======================================================
heading_paragraph(doc, 'Вывод по ходу работы')

body_paragraph(doc,
    'В ходе выполнения практической работы был изучен язык разметки '
    'гипертекста HTML и разработан многостраничный Web-сайт на тему '
    '«Каталог смартфонов». Сайт включает главную страницу, страницу каталога '
    'и три страницы отдельных товаров.')

body_paragraph(doc,
    'В процессе работы были практически освоены основные теги HTML: '
    'теги заголовков (<h1>–<h3>), параграфов (<p>), горизонтальной черты '
    '(<hr>), маркированных списков (<ul>, <li>), гиперссылок (<a>), '
    'изображений (<img>). Реализовано навигационное меню, единое для всех '
    'страниц, и подвал с информацией об авторских правах.')

body_paragraph(doc,
    'Приобретены навыки работы с относительными ссылками для перемещения '
    'между страницами сайта, а также с атрибутом target="_blank", '
    'позволяющим открывать содержимое в новой вкладке браузера. '
    'Цель работы достигнута в полном объёме.')

# ======================================================
#  СОХРАНЕНИЕ
# ======================================================
out = ('C:/Users/Vadim/PycharmProjects/клиент-серверные приложения/лр1/'
       'ПР-1_БИВТ-24-2_Горбачёв_В_А_Верстка_сайта_с_HTML.docx')
doc.save(out)
print('Saved:', out)
