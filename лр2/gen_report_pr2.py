from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ---- Параметры страницы: все поля 2 см ----
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(2)

# ---- Вспомогательные функции ----

def add_paragraph(document, text='', bold=False, italic=False, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0,
                  first_line_indent=None, space_before=0, space_after=0):
    p = document.add_paragraph()
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if first_line_indent is not None:
        pf.first_line_indent = Cm(first_line_indent)
    pPr = p._p.get_or_add_pPr()
    lnSpc = OxmlElement('w:spacing')
    if spacing == 1.0:
        lnSpc.set(qn('w:line'), '240')
        lnSpc.set(qn('w:lineRule'), 'auto')
    elif spacing == 1.5:
        lnSpc.set(qn('w:line'), '360')
        lnSpc.set(qn('w:lineRule'), 'auto')
    pPr.append(lnSpc)

    if text:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(size)
        r = run._r
        rPr = r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'),    'Times New Roman')
        rFonts.set(qn('w:hAnsi'),    'Times New Roman')
        rFonts.set(qn('w:cs'),       'Times New Roman')
        rPr.insert(0, rFonts)
    return p

def add_run_tnr(paragraph, text, bold=False, italic=False, size=14):
    run = paragraph.add_run(text)
    run.bold   = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),  'Times New Roman')
    rFonts.set(qn('w:hAnsi'),  'Times New Roman')
    rFonts.set(qn('w:cs'),     'Times New Roman')
    rPr.insert(0, rFonts)
    return run

def set_line_spacing(paragraph, spacing=1.5):
    pPr = paragraph._p.get_or_add_pPr()
    lnSpc = OxmlElement('w:spacing')
    if spacing == 1.5:
        lnSpc.set(qn('w:line'), '360')
    else:
        lnSpc.set(qn('w:line'), '240')
    lnSpc.set(qn('w:lineRule'), 'auto')
    pPr.append(lnSpc)

def body_paragraph(document, text, first_indent=True):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    if first_indent:
        pf.first_line_indent = Cm(1.25)
    else:
        pf.first_line_indent = Cm(0)
    set_line_spacing(p, 1.5)
    add_run_tnr(p, text)
    return p

def heading_paragraph(document, text, level=1):
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(6)
    pf.space_after  = Pt(6)
    pf.first_line_indent = Cm(0)
    set_line_spacing(p, 1.5)
    add_run_tnr(p, text, bold=True)
    return p

def add_page_break(document):
    p = document.add_paragraph()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p._p.append(br)

def code_paragraph(document, text):
    """Блок кода: Courier New 11pt, одинарный интервал, без отступа"""
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    pf.first_line_indent = Cm(0)
    set_line_spacing(p, 1.0)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(11)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'),  'Courier New')
    rFonts.set(qn('w:hAnsi'),  'Courier New')
    rFonts.set(qn('w:cs'),     'Courier New')
    rPr.insert(0, rFonts)
    return p

# ======================================================
#  ТИТУЛЬНЫЙ ЛИСТ
# ======================================================

top_lines = [
    'МИНИСТЕРСТВО НАУКИ И ВЫСШЕГО ОБРАЗОВАНИЯ РОССИЙСКОЙ ФЕДЕРАЦИИ',
    'ФЕДЕРАЛЬНОЕ ГОСУДАРСТВЕННОЕ АВТОНОМНОЕ ОБРАЗОВАТЕЛЬНОЕ УЧРЕЖДЕНИЕ '
    'ВЫСШЕГО ОБРАЗОВАНИЯ «НАЦИОНАЛЬНЫЙ ИССЛЕДОВАТЕЛЬСКИЙ ТЕХНОЛОГИЧЕСКИЙ '
    'УНИВЕРСИТЕТ «МИСИС»',
    'ИНСТИТУТ КОМПЬЮТЕРНЫХ НАУК (ИКН)',
    'КАФЕДРА АВТОМАТИЗИРОВАННЫХ СИСТЕМ УПРАВЛЕНИЯ (АСУ)',
]
for line in top_lines:
    add_paragraph(doc, line, bold=True, size=14,
                  align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

for _ in range(6):
    add_paragraph(doc, '', size=14, spacing=1.0)

add_paragraph(doc,
    'Курс «Клиент-серверные приложения и сетевые технологии»',
    bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

add_paragraph(doc, 'Практическая работа №\u00a02',
    bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

add_paragraph(doc, '«Основы каскадных таблиц стилей CSS»',
    bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

for _ in range(5):
    add_paragraph(doc, '', size=14, spacing=1.0)

def right_p(text, underline=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.paragraph_format.first_line_indent = Cm(0)
    set_line_spacing(p, 1.0)
    run = p.add_run(text)
    run.bold      = bold
    run.underline = underline
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    return p

right_p('Выполнил: студент группы БИВТ-24-2')
right_p('Горбачёв В.\u00a0А.', underline=True)
right_p('Проверили:           Абросимов Н.\u00a0А.')
right_p('                            Шелудяков П.\u00a0А.')

for _ in range(4):
    add_paragraph(doc, '', size=14, spacing=1.0)

add_paragraph(doc, 'Москва, 2026',
    bold=False, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, spacing=1.0)

add_page_break(doc)

# ======================================================
#  ЦЕЛЬ РАБОТЫ
# ======================================================
heading_paragraph(doc, 'Цель работы')

body_paragraph(doc,
    'Изучить и отработать на практике различные стили при создании '
    'страниц сайта с использованием каскадных таблиц стилей CSS. '
    'Освоить селекторы атрибутов и основные свойства стилей.')

# ======================================================
#  ХОД РАБОТЫ
# ======================================================
heading_paragraph(doc, 'Ход работы')

body_paragraph(doc,
    'Для выполнения работы использовался текстовый редактор '
    'Visual Studio Code и браузер Google Chrome. В качестве основы '
    'был взят многостраничный сайт «Каталог смартфонов», разработанный '
    'в практической работе №\u00a01.')

# --- Шаг 1 ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
p.paragraph_format.space_before = Pt(6)
set_line_spacing(p, 1.5)
add_run_tnr(p, 'Шаг 1. Создание внешнего файла стилей style.css и подключение к страницам.', bold=True)

body_paragraph(doc,
    'Был создан файл style.css, содержащий все стили сайта. '
    'Файл подключён ко всем HTML-страницам с помощью тега <link> '
    'в секции <head>:')

code_paragraph(doc, '<link rel="stylesheet" type="text/css" href="style.css">')

# --- Шаг 2 ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
p.paragraph_format.space_before = Pt(6)
set_line_spacing(p, 1.5)
add_run_tnr(p, 'Шаг 2. Установка фона страниц.', bold=True)

body_paragraph(doc,
    'Для тега body задан фоновый цвет в палитре RGB с помощью '
    'свойства background-color:')

code_paragraph(doc, 'body {')
code_paragraph(doc, '    background-color: rgb(240, 244, 248);')
code_paragraph(doc, '}')

# --- Шаг 3 ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
p.paragraph_format.space_before = Pt(6)
set_line_spacing(p, 1.5)
add_run_tnr(p, 'Шаг 3. Стилизация навигационного меню.', bold=True)

body_paragraph(doc,
    'Для меню сайта использован селектор класса .menu. Маркеры списка '
    'убраны с помощью свойства list-style: none. Для ссылок меню задан '
    'цвет текста rgb(0, 102, 153), размер шрифта 16px, полужирное '
    'начертание (font-weight: bold). При наведении курсора ссылка '
    'подчёркивается. Элементы меню расположены горизонтально с помощью '
    'display: inline-block.')

code_paragraph(doc, 'ul.menu {')
code_paragraph(doc, '    list-style: none;')
code_paragraph(doc, '    padding: 0;')
code_paragraph(doc, '}')
code_paragraph(doc, '')
code_paragraph(doc, 'ul.menu li a {')
code_paragraph(doc, '    color: rgb(0, 102, 153);')
code_paragraph(doc, '    font-size: 16px;')
code_paragraph(doc, '    font-weight: bold;')
code_paragraph(doc, '}')

# --- Шаг 4 ---
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.first_line_indent = Cm(1.25)
p.paragraph_format.space_before = Pt(6)
set_line_spacing(p, 1.5)
add_run_tnr(p, 'Шаг 4. Стилизация страницы товара.', bold=True)

body_paragraph(doc,
    'На каждой странице товара выделены три раздела: «Краткое описание '
    'товара», «Характеристики» и «Подробное описание товара». Для каждого '
    'раздела заданы стили согласно заданию.')

body_paragraph(doc,
    'Заголовки разделов (h3 с классом section-heading) оформлены '
    'следующим образом: чёрный цвет текста, размер шрифта 18px, '
    'насыщенность шрифта 400, фон в палитре RGB — rgb(200, 220, 240):')

code_paragraph(doc, 'h3.section-heading {')
code_paragraph(doc, '    color: black;')
code_paragraph(doc, '    font-size: 18px;')
code_paragraph(doc, '    font-weight: 400;')
code_paragraph(doc, '    background-color: rgb(200, 220, 240);')
code_paragraph(doc, '}')

body_paragraph(doc,
    'Для краткого описания товара (класс short-desc) задан цвет текста '
    'rgb(80, 80, 120), размер шрифта 14px, курсивное начертание '
    '(font-style: italic) и высота строки 16px (line-height):')

code_paragraph(doc, 'p.short-desc {')
code_paragraph(doc, '    color: rgb(80, 80, 120);')
code_paragraph(doc, '    font-size: 14px;')
code_paragraph(doc, '    font-style: italic;')
code_paragraph(doc, '    line-height: 16px;')
code_paragraph(doc, '}')

body_paragraph(doc,
    'Для подробного описания товара (класс full-desc) задан цвет текста '
    'rgb(60, 60, 80), размер шрифта 16px, насыщенность шрифта 400, '
    'высота строки 24px и выравнивание по левому краю (text-align: left):')

code_paragraph(doc, 'p.full-desc {')
code_paragraph(doc, '    color: rgb(60, 60, 80);')
code_paragraph(doc, '    font-size: 16px;')
code_paragraph(doc, '    font-weight: 400;')
code_paragraph(doc, '    line-height: 24px;')
code_paragraph(doc, '    text-align: left;')
code_paragraph(doc, '}')

body_paragraph(doc,
    'Для списка характеристик (класс specs) заданы отличные от '
    'основного текста стили: цвет rgb(30, 80, 100), размер шрифта 15px, '
    'увеличенная высота строки 28px. В качестве маркеров списка '
    'установлено произвольное изображение (marker.png) с помощью '
    'свойства list-style-image:')

code_paragraph(doc, 'ul.specs {')
code_paragraph(doc, "    list-style-image: url('images/marker.png');")
code_paragraph(doc, '    color: rgb(30, 80, 100);')
code_paragraph(doc, '    font-size: 15px;')
code_paragraph(doc, '    line-height: 28px;')
code_paragraph(doc, '}')

body_paragraph(doc,
    'На рисунке\u00a01 представлен внешний вид страницы товара iPhone 15 Pro '
    'после применения стилей CSS.')

p_fig = doc.add_paragraph()
p_fig.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_fig, 1.0)
add_run_tnr(p_fig, '[Рисунок 1 — Страница товара item_iphone.html со стилями CSS]', italic=True)

p_cap = doc.add_paragraph()
p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_cap, 1.0)
add_run_tnr(p_cap, 'Рисунок 1 — Страница товара со стилями CSS')

body_paragraph(doc,
    'На рисунке\u00a02 показана главная страница сайта с применёнными стилями.')

p_fig2 = doc.add_paragraph()
p_fig2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_fig2.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_fig2, 1.0)
add_run_tnr(p_fig2, '[Рисунок 2 — Главная страница index.html со стилями CSS]', italic=True)

p_cap2 = doc.add_paragraph()
p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_cap2.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_cap2, 1.0)
add_run_tnr(p_cap2, 'Рисунок 2 — Главная страница со стилями CSS')

# --- Таблица CSS-свойств ---
body_paragraph(doc,
    'В таблице\u00a01 приведены основные CSS-свойства, использованные в работе.')

p_tbl_cap = doc.add_paragraph()
p_tbl_cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_tbl_cap.paragraph_format.first_line_indent = Cm(0)
set_line_spacing(p_tbl_cap, 1.0)
add_run_tnr(p_tbl_cap, 'Таблица 1 — Основные CSS-свойства, использованные в работе')

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'

hdr = tbl.rows[0].cells
hdr[0].paragraphs[0].clear()
hdr[1].paragraphs[0].clear()
add_run_tnr(hdr[0].paragraphs[0], 'Свойство', bold=True)
add_run_tnr(hdr[1].paragraphs[0], 'Назначение', bold=True)
hdr[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
hdr[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

props = [
    ('background-color',  'Цвет фона элемента'),
    ('color',             'Цвет текста'),
    ('font-size',         'Размер шрифта'),
    ('font-weight',       'Насыщенность (жирность) шрифта'),
    ('font-style',        'Начертание шрифта (курсив, наклонный)'),
    ('line-height',       'Высота строки (межстрочный интервал)'),
    ('text-align',        'Горизонтальное выравнивание текста'),
    ('list-style',        'Стиль маркеров списка'),
    ('list-style-image',  'Произвольное изображение в качестве маркера списка'),
    ('text-decoration',   'Оформление текста (подчёркивание и др.)'),
    ('display',           'Способ отображения элемента (inline-block и др.)'),
]
for prop, desc in props:
    row = tbl.add_row()
    row.cells[0].paragraphs[0].clear()
    row.cells[1].paragraphs[0].clear()
    add_run_tnr(row.cells[0].paragraphs[0], prop)
    add_run_tnr(row.cells[1].paragraphs[0], desc)

# ======================================================
#  ВЫВОД
# ======================================================
heading_paragraph(doc, 'Вывод по ходу работы')

body_paragraph(doc,
    'В ходе выполнения практической работы были изучены основы '
    'каскадных таблиц стилей CSS. Создан внешний файл стилей style.css '
    'и подключён ко всем страницам сайта «Каталог смартфонов».')

body_paragraph(doc,
    'Были практически освоены следующие возможности CSS: задание фона '
    'страницы в палитре RGB; стилизация навигационного меню с удалением '
    'маркеров и горизонтальным расположением элементов; оформление '
    'заголовков с заданием цвета, размера и насыщенности шрифта; '
    'настройка параметров текста (размер, цвет, начертание, высота '
    'строки, выравнивание); использование произвольных изображений '
    'в качестве маркеров списка.')

body_paragraph(doc,
    'Были изучены три типа селекторов CSS: селекторы тегов, '
    'селекторы идентификаторов (id) и селекторы классов (class). '
    'В работе преимущественно использовались селекторы классов, '
    'что позволяет гибко применять стили к нужным элементам. '
    'Цель работы достигнута в полном объёме.')

# ======================================================
#  СОХРАНЕНИЕ
# ======================================================
import os
out_dir = os.path.dirname(os.path.abspath(__file__))
out = os.path.join(out_dir,
    'ПР-2_БИВТ-24-2_Горбачёв_В_А_Основы_CSS.docx')
doc.save(out)
print('Saved:', out)
