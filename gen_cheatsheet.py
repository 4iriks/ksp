from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    section.left_margin   = Cm(2)
    section.right_margin  = Cm(1.5)

# ---- Утилиты ----

def add_run(paragraph, text, bold=False, italic=False, size=12, font='Times New Roman'):
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = font
    run.font.size = Pt(size)
    r = run._r
    rPr = r.get_or_add_rPr()
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), font)
    rFonts.set(qn('w:hAnsi'), font)
    rFonts.set(qn('w:cs'), font)
    rPr.insert(0, rFonts)
    return run

def set_spacing(p, spacing=1.15):
    pPr = p._p.get_or_add_pPr()
    s = OxmlElement('w:spacing')
    val = str(int(spacing * 240))
    s.set(qn('w:line'), val)
    s.set(qn('w:lineRule'), 'auto')
    pPr.append(s)

def title(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(12)
    pf.space_after = Pt(6)
    set_spacing(p, 1.15)
    add_run(p, text, bold=True, size=18)
    return p

def h1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(14)
    pf.space_after = Pt(4)
    set_spacing(p, 1.15)
    add_run(p, text, bold=True, size=15)
    return p

def h2(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(8)
    pf.space_after = Pt(3)
    set_spacing(p, 1.15)
    add_run(p, text, bold=True, size=13)
    return p

def body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(1)
    pf.space_after = Pt(1)
    set_spacing(p, 1.15)
    add_run(p, text, size=12)
    return p

def code(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.left_indent = Cm(0.5)
    set_spacing(p, 1.0)
    add_run(p, text, size=10, font='Courier New')
    return p

def add_table(headers, rows):
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    for i, h in enumerate(headers):
        tbl.rows[0].cells[i].paragraphs[0].clear()
        add_run(tbl.rows[0].cells[i].paragraphs[0], h, bold=True, size=11)
        tbl.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    for row_data in rows:
        row = tbl.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].paragraphs[0].clear()
            # Check if it's code-like
            if val.startswith('<') or val.startswith('.') or val.startswith('#') or '()' in val or '{}' in val or ':' in val and '{' not in val and len(val) < 40:
                add_run(row.cells[i].paragraphs[0], val, size=10, font='Courier New')
            else:
                add_run(row.cells[i].paragraphs[0], val, size=11)
    return tbl

def page_break():
    p = doc.add_paragraph()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    p._p.append(br)

# ================================================================
#   ТИТУЛ
# ================================================================
title('Шпаргалка для защиты ПР 1-3')
body('Курс: Клиент-серверные приложения и сетевые технологии')
body('Горбачёв В. А., БИВТ-24-2')
body('')

# ================================================================
#   ПР-1: HTML
# ================================================================
h1('ПР-1. HTML — Язык разметки гипертекста')

h2('Что такое HTML?')
body('HTML (HyperText Markup Language) — язык разметки для создания веб-страниц. '
     'Браузер читает HTML-файл и отображает его содержимое. HTML состоит из тегов, '
     'которые описывают структуру страницы.')

h2('Структура HTML-документа')
code('<!DOCTYPE html>              — тип документа (HTML5)')
code('<html lang="ru">             — корневой элемент, язык страницы')
code('  <head>                     — мета-информация (не видна)')
code('    <meta charset="UTF-8">   — кодировка (кириллица)')
code('    <title>Заголовок</title> — текст во вкладке браузера')
code('  </head>')
code('  <body>                     — видимое содержимое')
code('    ...')
code('  </body>')
code('</html>')

h2('Все HTML-теги проекта')
add_table(['Тег', 'Назначение', 'Пример использования'], [
    ['<h1>, <h2>, <h3>', 'Заголовки (уровень 1-3, h1 — самый крупный)', 'Название сайта, товара, раздела'],
    ['<p>', 'Параграф — блок текста', 'Описания товаров, приветствие'],
    ['<a href="...">', 'Гиперссылка — переход на другую страницу', 'Меню, ссылки на товары'],
    ['<img src="..." alt="...">', 'Изображение', 'Фото смартфонов'],
    ['<ul>', 'Маркированный (неупорядоченный) список', 'Меню, характеристики'],
    ['<li>', 'Элемент списка (внутри <ul>)', 'Каждый пункт меню/характеристик'],
    ['<hr>', 'Горизонтальная черта-разделитель', 'Между секциями страницы'],
    ['<small>', 'Уменьшенный текст', 'Копирайт в подвале'],
    ['<br>', 'Принудительный перенос строки', 'Под картинкой в каталоге'],
    ['<div>', 'Блочный контейнер (группировка)', 'Карточки товаров, корзина'],
    ['<span>', 'Строчный контейнер', 'Текст внутри корзины'],
    ['<button>', 'Кнопка', 'Добавить в корзину, Оплатить'],
    ['<select> + <option>', 'Выпадающий список', 'Фильтр по бренду'],
    ['<label>', 'Подпись к элементу формы', 'Подпись к фильтру'],
    ['<strong> / <b>', 'Жирный текст', 'Цена товара'],
    ['<script src="...">', 'Подключение JavaScript', 'Перед </body>'],
    ['<link rel="stylesheet">', 'Подключение CSS-файла', 'В <head>'],
    ['<!-- комментарий -->', 'Комментарий (не виден)', 'Пометки в коде'],
])

h2('Все атрибуты тегов')
add_table(['Атрибут', 'Тег', 'Что делает'], [
    ['href="url"', '<a>', 'Адрес ссылки (относительный или абсолютный путь)'],
    ['target="_blank"', '<a>', 'Открыть ссылку в новой вкладке'],
    ['src="путь"', '<img>, <script>', 'Путь к файлу (картинка или скрипт)'],
    ['alt="текст"', '<img>', 'Текст, если картинка не загрузилась'],
    ['width="300"', '<img>', 'Ширина изображения в пикселях'],
    ['height="300"', '<img>', 'Высота изображения в пикселях'],
    ['lang="ru"', '<html>', 'Язык документа (для поисковиков, экранных читалок)'],
    ['charset="UTF-8"', '<meta>', 'Кодировка символов (поддержка русских букв)'],
    ['rel="stylesheet"', '<link>', 'Тип подключаемого ресурса — таблица стилей'],
    ['type="text/css"', '<link>', 'MIME-тип файла CSS'],
    ['class="имя"', 'любой', 'Назначить класс (для CSS-стилей)'],
    ['id="имя"', 'любой', 'Уникальный идентификатор (для CSS/JS)'],
    ['data-name="..."', 'любой', 'Произвольный data-атрибут (доступен из JS)'],
    ['data-price="..."', 'любой', 'Произвольный data-атрибут для цены'],
    ['data-category="..."', 'любой', 'Произвольный data-атрибут для категории'],
    ['for="id"', '<label>', 'Связь подписи с элементом формы по id'],
    ['value="..."', '<option>', 'Значение варианта в выпадающем списке'],
])

h2('Структура сайта')
code('index.html           — главная страница (приветствие)')
code('catalog.html         — каталог (3 товара с картинками)')
code('item_iphone.html     — iPhone 15 Pro')
code('item_samsung.html    — Samsung Galaxy S24 Ultra')
code('item_pixel.html      — Google Pixel 8 Pro')
code('style.css            — все CSS-стили')
code('script.js            — весь JavaScript-код')
code('images/              — папка с картинками')
code('images/marker.png    — кастомный маркер для списка')

h2('Что такое относительные ссылки?')
body('href="catalog.html" — файл в той же папке. '
     'href="images/foto.jpg" — файл в подпапке images. '
     'Не нужно писать полный путь (C:/Users/...), достаточно указать путь '
     'относительно текущего файла.')

page_break()

# ================================================================
#   ПР-2: CSS
# ================================================================
h1('ПР-2. CSS — Каскадные таблицы стилей')

h2('Что такое CSS?')
body('CSS (Cascading Style Sheets) — язык описания внешнего вида HTML-документа: '
     'цвета, шрифты, отступы, размеры, расположение элементов. '
     'Без CSS страница выглядит как голый текст с картинками.')

h2('Подключение CSS (внешний файл)')
code('<head>')
code('  <link rel="stylesheet" type="text/css" href="style.css">')
code('</head>')
body('rel="stylesheet" — говорит браузеру, что это стили. '
     'href="style.css" — путь к файлу. Один файл стилей подключается '
     'ко всем страницам — это удобно, меняешь в одном месте — меняется везде.')

h2('Синтаксис CSS')
code('селектор {')
code('    свойство: значение;')
code('    свойство2: значение2;')
code('}')

h2('Типы селекторов')
add_table(['Тип', 'Синтаксис', 'Пример', 'Что выбирает'], [
    ['Тега', 'тег', 'body { }', 'Все элементы <body>'],
    ['Класса', '.имя', '.menu { }', 'Все с class="menu"'],
    ['ID', '#имя', '#cart-list { }', 'Один с id="cart-list"'],
    ['Составной', 'тег.класс', 'h3.section-heading', 'Только <h3> с этим классом'],
    ['Вложенный', 'A B', 'ul.menu li a', '<a> внутри <li> внутри ul.menu'],
    ['Псевдокласс', ':hover', '.btn:hover', 'При наведении курсора'],
])

body('Приоритет: id (#) > class (.) > тег. Если конфликт — побеждает более специфичный селектор.')

h2('Все CSS-свойства проекта')

h2('Фон и цвета')
add_table(['Свойство', 'Значение в проекте', 'Что делает'], [
    ['background-color', 'rgb(240, 244, 248)', 'Цвет фона элемента'],
    ['color', 'black / rgb(80,80,120) / #333', 'Цвет текста'],
])

h2('Шрифт и текст')
add_table(['Свойство', 'Значение', 'Что делает'], [
    ['font-family', 'Arial, sans-serif', 'Шрифт (Arial, запасной — sans-serif)'],
    ['font-size', '14px / 16px / 18px', 'Размер шрифта в пикселях'],
    ['font-weight', '400 / bold', 'Жирность (400=обычный, 700=bold=жирный)'],
    ['font-style', 'italic / normal', 'Курсив или обычный'],
    ['line-height', '16px / 24px / 28px', 'Высота строки (межстрочный интервал)'],
    ['text-align', 'left / center', 'Горизонтальное выравнивание текста'],
    ['text-decoration', 'none / underline', 'Убрать/добавить подчёркивание'],
])

h2('Списки')
add_table(['Свойство', 'Значение', 'Что делает'], [
    ['list-style', 'none', 'Убрать стандартные маркеры (точки)'],
    ['list-style-image', "url('images/marker.png')", 'Своя картинка вместо маркера'],
])

h2('Размеры и отступы (box model)')
body('Каждый элемент — это коробка: content (содержимое) → padding (внутренний отступ) → '
     'border (рамка) → margin (внешний отступ).')
add_table(['Свойство', 'Значение', 'Что делает'], [
    ['width', '200px', 'Ширина элемента'],
    ['margin', '20px / 10px 15px', 'Внешний отступ (от соседей)'],
    ['padding', '10px / 6px 12px', 'Внутренний отступ (от содержимого до рамки)'],
    ['margin-right', '20px', 'Внешний отступ только справа'],
    ['margin-top', '20px', 'Внешний отступ только сверху'],
])
body('Запись "10px 15px" — сверху/снизу 10px, слева/справа 15px. '
     'Запись "6px 12px" — аналогично.')

h2('Расположение элементов')
add_table(['Свойство', 'Значение', 'Что делает'], [
    ['display', 'inline-block', 'Блок, но в строку (карточки товаров, пункты меню)'],
    ['display', 'flex', 'Flexbox — гибкое расположение дочерних элементов'],
    ['display', 'none', 'Скрыть элемент полностью'],
    ['vertical-align', 'top', 'Выровнять по верхнему краю'],
    ['justify-content', 'space-between', 'Flexbox: растянуть по горизонтали с равным промежутком'],
    ['align-items', 'center', 'Flexbox: выровнять по центру по вертикали'],
])

h2('Оформление')
add_table(['Свойство', 'Значение', 'Что делает'], [
    ['border', '1px solid #ddd', 'Рамка: 1px, сплошная, серая'],
    ['border-radius', '8px', 'Скругление углов'],
    ['border-bottom', '1px solid #eee', 'Рамка только снизу'],
    ['cursor', 'pointer', 'Курсор-рука при наведении (как у ссылки)'],
])

h2('3 способа задать цвет')
code('color: red;                /* По имени (ограниченный набор) */')
code('color: #ff0000;            /* HEX — шестнадцатеричный (#RRGGBB) */')
code('color: rgb(255, 0, 0);     /* RGB — красный, зелёный, синий (0-255) */')
body('rgb() — самый гибкий. Каждый канал от 0 до 255. '
     'rgb(0,0,0) = чёрный, rgb(255,255,255) = белый.')

h2('Классы в HTML и их CSS-селекторы')
add_table(['HTML', 'CSS-селектор', 'Для чего'], [
    ['class="menu"', 'ul.menu / .menu', 'Навигационное меню'],
    ['class="section-heading"', 'h3.section-heading', 'Заголовки разделов товара'],
    ['class="short-desc"', 'p.short-desc', 'Краткое описание (курсив, 14px)'],
    ['class="full-desc"', 'p.full-desc', 'Подробное описание (16px, line-height 24px)'],
    ['class="specs"', 'ul.specs', 'Список характеристик (кастомный маркер)'],
    ['class="footer"', '.footer', 'Подвал страницы'],
    ['class="product-card"', '.product-card', 'Карточка товара в каталоге'],
    ['class="filter-bar"', '.filter-bar', 'Панель фильтра'],
    ['class="cart-section"', '.cart-section', 'Блок корзины'],
    ['class="btn btn-add"', '.btn / .btn-add', 'Кнопка (общие стили + цвет)'],
])

page_break()

# ================================================================
#   ПР-3: JavaScript
# ================================================================
h1('ПР-3. JavaScript — Добавление динамики')

h2('Что такое JavaScript?')
body('JavaScript (JS) — язык программирования, который выполняется в браузере. '
     'Позволяет делать страницу интерактивной: реагировать на клики, '
     'менять содержимое, показывать/скрывать элементы без перезагрузки страницы.')

h2('Подключение JS')
code('<script src="script.js"></script>   <!-- Перед </body> -->')
body('Почему перед </body>? Чтобы весь HTML уже был загружен и JS мог найти элементы на странице. '
     'Если подключить в <head>, элементов ещё нет — скрипт не найдёт их.')

h2('Что такое DOM?')
body('DOM (Document Object Model) — представление HTML-документа в виде дерева объектов. '
     'Каждый HTML-тег становится объектом, который JS может найти, '
     'изменить, удалить или создать заново. document — корень этого дерева.')

h2('Переменные: let vs const')
add_table(['Ключевое слово', 'Можно менять?', 'Пример', 'Когда использовать'], [
    ['let', 'Да', 'let cart = [];  cart = [1,2]', 'Значение будет меняться (корзина, счётчик)'],
    ['const', 'Нет', 'const products = [...]', 'Значение не меняется (массив товаров, ссылки на DOM)'],
])
body('Важно: const для массива означает, что нельзя присвоить НОВЫЙ массив, '
     'но можно менять содержимое (push, splice). let cart = [] — тут нужен let, '
     'потому что мы делаем cart = [] при очистке (присваиваем новый массив).')

h2('Типы данных')
add_table(['Тип', 'Пример', 'Описание'], [
    ['Число (number)', '129990, 0, 3.14', 'Целые и дробные числа'],
    ['Строка (string)', '"iPhone", "Корзина пуста!"', 'Текст в кавычках'],
    ['Булевый (boolean)', 'true, false', 'Истина / ложь'],
    ['Массив (array)', '[1, 2, 3], []', 'Упорядоченный список значений'],
    ['Объект (object)', '{ name: "iPhone", price: 129990 }', 'Набор пар ключ: значение'],
])

h2('Массив — список значений')
code('let cart = [];                          // Пустой массив')
code('cart.push({ name: "iPhone" });          // Добавить в конец')
code('cart.splice(2, 1);                      // Удалить 1 элемент с индекса 2')
code('cart.length;                            // Кол-во элементов (0 = пустой)')
code('cart = [];                              // Очистить')
code('cart.forEach(item => { ... });          // Перебрать все элементы')

h2('Объект — набор "ключ: значение"')
code('const product = {')
code('    id: 1,')
code('    name: "iPhone 15 Pro",')
code('    price: 129990,')
code('    category: "apple"')
code('};')
code('product.name     // → "iPhone 15 Pro"')
code('product.price    // → 129990')
code('// Краткая запись: { name, price } = { name: name, price: price }')

h2('Функции')
body('Функция — блок кода, который можно вызывать многократно.')
code('// Обычная функция:')
code('function sum(a, b) {')
code('    return a + b;')
code('}')
code('')
code('// Стрелочная функция (arrow function) — короткая запись:')
code('const sum = (a, b) => a + b;')
code('')
code('// Если тело в несколько строк — нужны {} и return:')
code('const calculateTotal = () => {')
code('    let total = 0;')
code('    cart.forEach(item => total += item.price);')
code('    return total;')
code('};')
code('')
code('// Один параметр — скобки можно опустить:')
code('const double = x => x * 2;')

h2('Все функции проекта (script.js)')
add_table(['Функция', 'Тип', 'Что делает'], [
    ['calculateTotal()', 'Стрелочная', 'Считает сумму цен в корзине через forEach'],
    ['addToCart(product)', 'Стрелочная', 'Добавляет товар в cart (push), перерисовывает'],
    ['removeFromCart(index)', 'Стрелочная', 'Удаляет товар по индексу (splice), перерисовывает'],
    ['clearCart()', 'Стрелочная', 'Обнуляет cart = [], перерисовывает'],
    ['pay()', 'Стрелочная', 'Пустая корзина → alert; иначе → alert с суммой + очистка'],
    ['renderCart()', 'Стрелочная', 'Очищает HTML-список, рисует заново, обновляет сумму'],
    ['filterProducts(category)', 'Стрелочная', 'Перебирает карточки, скрывает/показывает по категории'],
])

h2('Получение элементов (DOM-запросы)')
code('document.querySelector("#cart-list")         // 1 элемент по id')
code('document.querySelector(".btn")               // 1-й элемент с классом')
code('document.querySelectorAll(".product-card")   // ВСЕ с классом (NodeList)')
body('querySelector принимает CSS-селектор. "#id" — по id, ".class" — по классу, '
     '"tag" — по тегу. querySelectorAll возвращает список всех совпадений.')

h2('События (addEventListener)')
code('element.addEventListener("событие", функция);')
code('')
code('// Примеры:')
code('button.addEventListener("click", () => { ... });       // Клик')
code('select.addEventListener("change", () => { ... });      // Изменение выбора')
code('document.addEventListener("DOMContentLoaded", () => {  // HTML загружен')
code('    // инициализация')
code('});')

add_table(['Событие', 'Когда срабатывает', 'Где используется'], [
    ['"click"', 'Клик мышкой', 'Кнопки: добавить, оплатить, удалить, очистить'],
    ['"change"', 'Изменение значения элемента', 'Фильтр (select)'],
    ['"DOMContentLoaded"', 'Весь HTML загружен', 'Инициализация обработчиков при старте'],
])

h2('data-атрибуты (хранение данных в HTML)')
code('<!-- HTML -->')
code('<div class="product-card" data-category="apple">')
code('<button data-name="iPhone 15 Pro" data-price="129990">')
code('')
code('// JavaScript')
code('card.dataset.category      // → "apple"')
code('button.dataset.name        // → "iPhone 15 Pro"')
code('button.dataset.price       // → "129990" (СТРОКА!)')
code('Number(button.dataset.price) // → 129990 (число)')
body('dataset всегда возвращает строки. Для математики нужно Number() '
     'или parseInt() / parseFloat().')

h2('Управление элементами')
code('// Скрыть/показать:')
code('element.style.display = "none";         // Скрыть')
code('element.style.display = "inline-block"; // Показать')
code('')
code('// Изменить текст:')
code('element.textContent = "Итого: 500 руб.";')
code('')
code('// Очистить HTML внутри:')
code('element.innerHTML = "";')
code('')
code('// Задать класс:')
code('element.className = "cart-item";')

h2('Создание элементов (динамический HTML)')
code('const div = document.createElement("div");    // Создать <div>')
code('const span = document.createElement("span");  // Создать <span>')
code('span.textContent = "Текст";                   // Задать текст')
code('div.appendChild(span);                         // Вставить span в div')
code('cartList.appendChild(div);                     // Вставить div на страницу')

h2('alert — всплывающее окно')
code('alert("Корзина пуста!");')
code('alert("Оплата прошла успешно! Сумма: " + total + " руб.");')
body('alert() блокирует страницу до нажатия ОК. Используется для простых уведомлений.')

h2('forEach — перебор массива')
code('cart.forEach((item, index) => {')
code('    // item  — текущий элемент массива')
code('    // index — его порядковый номер (0, 1, 2...)')
code('    total += item.price;')
code('});')
body('forEach вызывает функцию для каждого элемента массива. '
     'Первый аргумент — сам элемент, второй (необязательный) — его индекс.')

page_break()

# ================================================================
#   АЛГОРИТМ РАБОТЫ (как всё связано)
# ================================================================
h1('Как всё работает вместе (алгоритм)')

h2('Фильтр товаров')
body('1. Пользователь выбирает бренд в <select id="category-filter">')
body('2. Срабатывает событие "change"')
body('3. JS получает выбранное значение: filterSelect.value (например "apple")')
body('4. Функция filterProducts() перебирает все .product-card')
body('5. Сравнивает card.dataset.category с выбранным значением')
body('6. Совпало → display = "inline-block" (показать)')
body('7. Не совпало → display = "none" (скрыть)')
body('8. Если выбрано "all" — показываются все')

h2('Добавление в корзину')
body('1. Пользователь кликает "Добавить в корзину"')
body('2. Срабатывает событие "click" на кнопке')
body('3. JS читает data-name и data-price из кнопки')
body('4. Цена преобразуется из строки в число: Number(price)')
body('5. Создаётся объект { name, price } и передаётся в addToCart()')
body('6. addToCart() делает cart.push(product) и вызывает renderCart()')
body('7. renderCart() очищает список (innerHTML = ""), рисует заново каждый товар')
body('8. Пересчитывает сумму через calculateTotal() и обновляет textContent')

h2('Оплата')
body('1. Пользователь кликает "Оплатить"')
body('2. Функция pay() проверяет cart.length === 0')
body('3. Если пусто → alert("Корзина пуста!")')
body('4. Если не пусто → alert("Оплата прошла успешно! Сумма: ...") → clearCart()')

h2('Удаление товара')
body('1. Пользователь кликает "Удалить" у конкретного товара')
body('2. removeFromCart(index) — удаляет элемент по индексу через splice')
body('3. Перерисовка корзины (renderCart)')

page_break()

# ================================================================
#   ВОЗМОЖНЫЕ ВОПРОСЫ НА ЗАЩИТЕ
# ================================================================
h1('Возможные вопросы на защите')

questions = [
    ('Что такое HTML?',
     'Язык разметки для создания веб-страниц. Состоит из тегов, описывающих структуру документа.'),
    ('Что делает <!DOCTYPE html>?',
     'Указывает браузеру, что это документ HTML5. Без него браузер может неправильно отобразить страницу.'),
    ('Зачем charset="UTF-8"?',
     'Задаёт кодировку. Без неё русские буквы могут отображаться как кракозябры.'),
    ('Разница между <div> и <span>?',
     '<div> — блочный (занимает всю ширину, новая строка). <span> — строчный (внутри текста, не разрывает строку).'),
    ('Что такое CSS?',
     'Каскадные таблицы стилей — язык описания внешнего вида HTML. Цвета, шрифты, отступы, расположение.'),
    ('Почему внешний CSS-файл лучше встроенного?',
     'Один файл подключается ко всем страницам. Изменил в одном месте — изменилось везде. Разделение контента и оформления.'),
    ('Что такое каскадность в CSS?',
     'Если несколько правил применяются к одному элементу, побеждает более специфичное. id > class > тег.'),
    ('Что такое box model?',
     'Каждый элемент — коробка: content → padding → border → margin. Padding — внутренний отступ, margin — внешний.'),
    ('Как подключается CSS?',
     '<link rel="stylesheet" type="text/css" href="style.css"> в <head>.'),
    ('Как подключается JS?',
     '<script src="script.js"></script> перед </body>.'),
    ('Почему скрипт перед </body>?',
     'Чтобы HTML уже загрузился и JS мог найти элементы. Иначе querySelector вернёт null.'),
    ('Что такое DOM?',
     'Document Object Model — представление HTML в виде дерева объектов. JS работает с DOM, чтобы менять страницу.'),
    ('Разница let и const?',
     'let — можно менять значение позже. const — нельзя (константа). Для массивов const разрешает push/splice, но не переприсваивание.'),
    ('Что такое стрелочная функция?',
     'Короткая запись: const fn = (x) => x * 2. Вместо function fn(x) { return x * 2; }.'),
    ('Как работает фильтр?',
     'Берём value из <select>, перебираем карточки querySelectorAll, сравниваем dataset.category, скрываем/показываем через display.'),
    ('Как работает корзина?',
     'Глобальный массив cart. push добавляет, splice удаляет, forEach перебирает для подсчёта суммы. renderCart перерисовывает HTML.'),
    ('Что такое data-атрибуты?',
     'Атрибуты data-* в HTML для хранения данных. Доступны через element.dataset в JS. Всегда строки — нужен Number() для чисел.'),
    ('Зачем Number()?',
     'dataset возвращает строки. "129990" + "84990" = "12999084990" (склеивание). Number("129990") + Number("84990") = 214980 (сложение).'),
    ('Что делает addEventListener?',
     'Назначает функцию-обработчик на событие. element.addEventListener("click", fn) — при клике вызовется fn.'),
    ('Что такое innerHTML и textContent?',
     'textContent — только текст. innerHTML — HTML-разметка. innerHTML = "" очищает содержимое элемента.'),
    ('Что делает splice?',
     'Удаляет/вставляет элементы массива. cart.splice(2, 1) — удаляет 1 элемент с индекса 2.'),
    ('Что делает push?',
     'Добавляет элемент в конец массива. cart.push(product).'),
    ('Что такое createElement?',
     'Создаёт новый HTML-элемент в памяти. document.createElement("div") → <div>. Потом appendChild вставляет его на страницу.'),
]

for q, a in questions:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.space_before = Pt(4)
    pf.space_after = Pt(2)
    set_spacing(p, 1.15)
    add_run(p, q, bold=True, size=12)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf2 = p2.paragraph_format
    pf2.space_before = Pt(0)
    pf2.space_after = Pt(4)
    pf2.left_indent = Cm(0.5)
    set_spacing(p2, 1.15)
    add_run(p2, a, size=12)

# ================================================================
import os
out_dir = os.path.dirname(os.path.abspath(__file__))
out = os.path.join(out_dir, 'Шпаргалка_ПР-1-3.docx')
doc.save(out)
print('Saved:', out)
